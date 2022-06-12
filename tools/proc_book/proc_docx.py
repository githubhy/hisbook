import os
import logging
import argparse
import itertools
from docx import Document
import pyparsing as pp
from pyparsing import pyparsing_unicode as ppu

parser = argparse.ArgumentParser(description='')
parser.add_argument('--output_dir', type=str, default=None)
parser.add_argument('file_name')
args = parser.parse_args()

log = logging.getLogger(__name__)
logging.basicConfig(level=logging.DEBUG,
    format="[%(asctime)s] %(levelname)s [%(name)s.%(funcName)s:%(lineno)d] %(message)s",
    datefmt="%d/%b/%Y %H:%M:%S")


doc = Document(args.file_name)

def add_comment(texts, level=0):
    if type(texts) is list:
        indent = ' ' * 4
        for text in texts:
            text = add_comment(text, level+1)
        if level > 0:
            texts.insert(0, '\n' + indent*level + '[')
            texts.append(']{{.cmt{0}}}\n{1}'.format(level, indent*(level-1)))
    return texts


def flatten(x):
    result = []
    for el in x:
        if hasattr(el, "__iter__") and not isinstance(el, str):
            result.extend(flatten(el))
        else:
            result.append(el)
    return result


def iter_text(paragraphs):
    '''
    Split paragraphs as this post says: https://stackoverflow.com/a/60424622
    '''

    postfix = ''

    for k1, g1 in itertools.groupby(paragraphs, lambda e: e.style.name.startswith('Heading 1')):
        if k1:
            yield '{0}{1}{2}'.format('# ', ''.join([p.text for p in g1]), postfix)
        else:
            for k2, g2 in itertools.groupby(g1, lambda e: e.style.name.startswith('Heading 2')):
                if k2:
                    yield '{0}{1}{2}'.format('## ', ''.join([p.text for p in g2]), postfix)
                else:
                    temp = ''
                    for k3, g3 in itertools.groupby(g2, lambda e: (e.paragraph_format.first_line_indent != None and
                                                                 int(e.paragraph_format.first_line_indent) > 0)
                                                                 or (e.paragraph_format.left_indent != None and
                                                                 int(e.paragraph_format.left_indent) > 0)):
                        text = ''.join([p.text for p in g3])
                        if k3:
                            temp = text
                        else:
                            yield '{2}{0}{1}'.format(text, postfix, temp)



def test_cmts_parser(parser):
    result = parser.run_tests('''
            # Empty string

            # Normal string
            八月，【江夏郡，见《武纪》【建安】十三年。】太守文聘坚守。【◎胡三省注：○文聘时屯石阳】

            # Medium string
            六月甲戌，任城王彰薨于京都。【◎《任城王传》注引《魏氏春秋》云：初，彰问玺绶，将有异志，故来朝不即得见。彰忿怒暴薨。◎《陈思王传》注引《魏氏春秋》云：是时待遇诸国法峻。任城王暴薨。诸王怀友于之痛。〖参阅《任城王传》注引《世说新语》。〗◎是任城王之死，实为魏文所害，当时陈王亦几不免。天性凉薄，宜其享国之不永也。】甲申，太尉贾诩薨。太白昼见。

            # Long string
            六月甲戌，任城王彰薨于京都。【◎《任城王传》注引《魏氏春秋》云：初，彰问玺绶，将有异志，故来朝不即得见。彰忿怒暴薨。◎《陈思王传》注引《魏氏春秋》云：是时待遇诸国法峻。任城王暴薨。诸王怀友于之痛。〖参阅《任城王传》注引《世说新语》。〗◎是任城王之死，实为魏文所害，当时陈王亦几不免。天性凉薄，宜其享国之不永也。】甲申，太尉贾诩薨。太白昼见。【◎《晋书·天文志下》云：黄初四年六月甲申，太白昼见。案刘向《五行纪论》曰：“太白少阴，弱，不得专行，故以己未为界，不得经天而行。经天则昼见，其占为兵丧，为不臣，为更王；强国弱，小国强。”是时孙权受魏爵号，而称兵距守。◎弼按：孙权背魏，在黄初三年，与四年之太白昼见何涉？以是知天文、五行、符瑞诸志，多不足据也。】

            # Extra long string
            【◎臣松之按：魏为土行，故服色尚黄。行殷之时，以建丑为正，故牺牲旂旗一用殷礼。《礼记》云：“夏后氏尚黑，故戎事乘骊，牲用玄；殷人尚白，戎事乘翰，牲用白；周人尚赤，戎事乘騵，牲用骍。”郑玄云：“夏后氏以建寅为正，物生色黑；殷以建丑为正，物牙色白；周以建子为正，物萌色赤。翰，白色马也。《易》曰‘白马翰如’。”《周礼·巾车职》“建大赤以朝”，“大白以即戎”，此则周以正色之旗以朝，先代之旗即戎。今魏用殷礼，变周之制，故建大白以朝，大赤即戎。改太和历曰景初历。【◎《宋书·律志序》云：窃以班氏《律历》，前事已详，自杨伟改创《景初》，而《魏书》阙志。及元嘉重造新法，大明博议回改。自魏至宋，宜入今书。◎又《历志（上）[中]》云：魏明帝景初元年，改定历数，以建丑之月为正，改其年三月为孟夏四月。其孟仲季月，虽与正岁不同，至于郊祀、迎气、祭祠、烝尝，巡狩、蒐田，分至启闭，班宣时令，皆以建寅为正。三年正月，帝崩，复用夏正。杨伟表曰：“臣揽载籍，断考历数，时以纪农，月以纪事，其所由来，遐而尚矣。乃自少昊，则玄鸟司分；颛顼、帝喾，则重、黎司天；唐帝、虞舜，则羲、和掌日。三代因之，则世有日官。日官司历，则颁之诸侯，诸侯受之，则颁于境内。夏后之代，羲、和湎淫，废时乱日，则《书》载《胤征》。由此观之，审农时而重人事者，历代然也。逮至周室既衰，战国横骛，告朔之羊，废而不绍，登台之礼，灭而不遵。闰分乖次而不识，孟陬失纪而莫悟，大火犹西流，而怪蛰虫之不藏也。是时也，天子不协时，司历不书日，诸侯不受职，日御不分朔，人事不恤，废弃农时。仲尼之拨乱于《春秋》，托褒贬纠正，司历失闰，则讥而书之，登台颁朔，则谓之有礼。自此以降，暨于秦、汉，乃复以孟冬为岁首，闰为后九月，中节乖错，时月纰缪，加时后天，蚀不在朔，累载相袭，久而不革也。至武帝元封七年，始乃寤其缪焉。于是改正朔，更历数，使大才通人，造《太初历》。校中朔所差，以正闰分；课中星得度，以考疏密，以建寅之月为正朔，以黄钟之月为历初。其历斗分太多，后遂疏阔。至元和二年，复用《四分历》。施而行之。至于今日，考察日蚀，率常在晦，是则斗分太多，故先密后疏而不可用也。是以臣前以制典余日，推考天路，稽之前典，验之食朔，详而精之，更建密历，则不先不后，古今中天。以昔在唐帝，协日正时，允厘百工，咸熙庶绩也。欲使当今国之典礼，凡百制度，皆韬合往古，郁然备足，乃改正朔，更历数，以大吕之月为岁首，以建子之月为历初。臣以为昔在帝代，则法曰《颛顼》，曩自轩辕，则历曰《黄帝》。暨至汉之孝武，革正朔，更历数，改元曰太初，因名《太初历》。今改元为景初，宜曰《景初历》。臣之所建《景初历》，法数则约要，施用则近密，治之则省功，学之则易知。虽复使研、桑心算，隶首运筹，重、黎司晷，羲、和察景，以考天路，步验日月，究极精微，尽术数之极者，皆未如臣如此之妙也。是以累代历数，皆疏而不密，自黄帝以来，改革不已。”◎何焯曰：景初历，尚书郎杨伟所造，事详《宋书·历志》。曹爽有参军杨伟，疑即斯人。《宋书》又载黄初中太史丞韩翊尝造《黄初历》，时陈群为尚书令，奏以为是非得失，当一年决定，今注家于《群传》遗之。杨伟历施用暨于晋、宋，而名字翳然，亦采掇之缺略也。◎弼按：杨伟，事见《高堂隆传》注引《魏略》云“诏使隆与尚书郎杨伟、太史待诏骆禄参共推校”，何氏谓“注家采掇缺略”，殆未细检耳。】其春夏秋冬孟仲季月虽与正岁不同，至于郊祀、迎气、礿祠、蒸尝、【◎《宋书·礼志》作“礿祀烝尝”，《历志》作“祭祀烝尝”。◎《尔雅·释诂》：禋、祀、祠、蒸、尝、禴，祭也。◎郭璞注：《书》曰“禋于六宗”，余者皆以为四时祭名也。◎陆德明《音义》：祠，春祭名；蒸，冬祭名；尝，秋祭名；禴，又作“礿”，夏祭名。◎《释天》：春祭曰祠，夏祭曰礿，秋祭曰尝，冬祭曰蒸。◎郭璞注：祠之言食也。礿者，新菜可汋。尝者，尝新谷也。蒸者，进品物也。◎《诗·小雅·天保》：禴祠烝尝。◎《毛传》云：春曰祠，夏曰禴，秋曰尝，冬曰烝。◎《礼记·王制》：天子诸侯宗庙之祭，春曰礿，夏曰禘，秋曰尝，冬曰烝。◎郑注云：○此盖夏、殷之祭名，周则改之，春曰祠，夏曰礿，以禘为殷祭。《诗·小雅》曰“礿祠烝尝，于公先王”。此周时祭宗庙之名。○孔《疏》引皇氏云：礿，薄也。春物未成，其祭品鲜薄也。○《诗·天保篇》谓文王受命，已改殷之祭名，以夏祭之禘，改名曰礿。而《诗》先言（灼）[礿]后祠者，从便文，尝在烝下，以韵句也。】巡狩、蒐田、分至启闭、班宣时令、中气早晚、敬授民事，皆以正岁斗建为历数之序。
            ''')
    res_list = result[1][0][1].as_list()
    print(res_list)
    print(''.join(flatten(res_list)))
    print(''.join(flatten(add_comment(res_list))))


def test_docx(pars):
    for par in pars:
        pf = par.paragraph_format
        if 1 or (pf.first_line_indent != None and int(pf.first_line_indent) > 0) or (pf.left_indent != None and int(pf.left_indent) > 0):
            print('{0} {4} -- {1} -- {3} -- {2}'.format(pf.first_line_indent, pf.alignment, par.text[0:10], pf.space_before, pf.left_indent))

try:
    file_head, file_extension = os.path.splitext(args.file_name)
    file_path_and_name = os.path.split(file_head)
except:
    log.error("No file name specified or invalid filename")
    exit(1)

if args.output_dir:
    os.makedirs(args.output_dir, exist_ok=True) 
    filename = os.path.join(args.output_dir, file_path_and_name[1])
else:
    filename = os.path.join(file_path_and_name[0], file_path_and_name[1])

out_filename = filename + '.Rmd'

content = pp.Combine(pp.Word(pp.printables + ppu.Chinese.printables + '〇 　。，；：、-！￥……（）——？《》〖〗·“”„’‘◎○,;:/-!?<>~()[]{}#@$%^&*+=_/\\|`\'\'"'), adjacent = False) # Combine() is space sensitive, see: https://stackoverflow.com/a/2940844
# content = pp.Combine(pp.Regex(r'[^【】]*')) # Regex() is very time consuming
comment = pp.nested_expr('【', '】', content=content)
cmts_parser = ((content + comment[...]) ^ comment)[...]
pp.enable_all_warnings()

# test_cmts_parser(cmts_parser); exit(0)
# test_docx(doc.paragraphs); exit(0)

# Parse nested comments as this post says: https://stackoverflow.com/a/5454510
styled_texts = []
FLAG_DEBUG = 0
for text in iter_text(doc.paragraphs):
    try:
        res = cmts_parser.parse_string(text, parse_all=True)
    except pp.ParseException as pe:
        print(pe.explain(depth=0)) # logging.warning
    styled_text = ''.join(flatten(add_comment(res.as_list())))
    styled_texts.append(styled_text)
    if FLAG_DEBUG:
        print('[original text]: ' + text)
        print('[parsed text]: ' + styled_text)

with open(out_filename, 'w') as f:
    f.write('\n\n'.join(styled_texts))
