import logging
import argparse
import itertools
from docx import Document
import pyparsing as pp
from pyparsing import pyparsing_unicode as ppu
from utils import FilenameInOut, flatten, MdIdAttacher


parser = argparse.ArgumentParser(description='')
parser.add_argument('--output_dir', type=str, default=None)
parser.add_argument('--test_cat', type=str, default=None)
parser.add_argument('file_name')
args = parser.parse_args()

log = logging.getLogger(__name__)
logging.basicConfig(level=logging.CRITICAL,
    format="[%(asctime)s] %(levelname)s [%(name)s.%(funcName)s:%(lineno)d] %(message)s",
    datefmt="%d/%b/%Y %H:%M:%S")


def cmt_ele(text_before, level):
    left = '[' #'{{.c{0}}}['.format(level) #'<span class="c{0}">'.format(level)
    right = ']{{.c{0}}}'.format(level) #'</span>\n'.format(level)
    cmt_pre_ele = ('\n' if isinstance(text_before, str) else '') + left
    cmt_post_ele = right + '\n'
    return (cmt_pre_ele, cmt_post_ele)

def add_comment(texts, text_before = [], level = 0):
    if type(texts) is list:
        for i in range(len(texts)):
            # When it is like 【*【 and 【str*【str】, a newline should be inserted in place of *.
            texts[i] = add_comment(texts[i], texts[i-1] if i > 0 else '', level+1)

        if level > 0:
            cmt_pre, cmt_post = cmt_ele(text_before, level)
            texts = [text.replace('[', '〈').replace(']', '〉') if isinstance(text, str) else text for text in texts]
            texts.insert(0, cmt_pre)
            texts.append(cmt_post)

    return texts


def iter_text(paragraphs):
    '''
    Split paragraphs as this post says: https://stackoverflow.com/a/60424622
    '''

    postfix = ''

    for k1, g1 in itertools.groupby(paragraphs, lambda e: e.style.name.startswith('Heading 1')):
        if k1:
            yield '{0}{1}{2}'.format('# ', ''.join([p.text for p in g1]), postfix)
        else:
            for k2, g2 in itertools.groupby(g1, lambda e: e.style.name.startswith('Heading 2')):
                if k2:
                    yield '{0}{1}{2}'.format('## ', ''.join([p.text for p in g2]), postfix)
                else:
                    temp = ''
                    for k3, g3 in itertools.groupby(
                                    g2, lambda e: (e.paragraph_format.first_line_indent != None and
                                                int(e.paragraph_format.first_line_indent) > 0)
                                                or (e.paragraph_format.left_indent != None and
                                                int(e.paragraph_format.left_indent) > 0 and
                                                # paragaph may starts with a bullet as a list item
                                                not e.style.name.startswith('List Paragraph'))
                                                ):
                        text = ''.join([p.text for p in g3])
                        if k3:
                            temp = text
                        else:
                            yield '{2}{0}{1}'.format(text, postfix, temp)



def test_cmts_parser(parser):
    result = parser.run_tests('''
            # Empty string

            # Normal string
            八月，【江夏郡，见《武纪》【建安】十三年。】太守文聘坚守。【◎胡三省注：○文聘时屯石阳】

            # Medium string
            六月甲戌，任城王彰薨于京都。【◎《任城王传》注引《魏氏春秋》云：初，彰问玺绶，将有异志，故来朝不即得见。彰忿怒暴薨。◎《陈思王传》注引《魏氏春秋》云：是时待遇诸国法峻。任城王暴薨。诸王怀友于之痛。〖参阅《任城王传》注引《世说新语》。〗◎是任城王之死，实为魏文所害，当时陈王亦几不免。天性凉薄，宜其享国之不永也。】甲申，太尉贾诩薨。太白昼见。

            # Long string
            六月甲戌，任城王彰薨于京都。【◎《任城王传》注引《魏氏春秋》云：初，彰问玺绶，将有异志，故来朝不即得见。彰忿怒暴薨。◎《陈思王传》注引《魏氏春秋》云：是时待遇诸国法峻。任城王暴薨。诸王怀友于之痛。〖参阅《任城王传》注引《世说新语》。〗◎是任城王之死，实为魏文所害，当时陈王亦几不免。天性凉薄，宜其享国之不永也。】甲申，太尉贾诩薨。太白昼见。【◎《晋书·天文志下》云：黄初四年六月甲申，太白昼见。案刘向《五行纪论》曰：“太白少阴，弱，不得专行，故以己未为界，不得经天而行。经天则昼见，其占为兵丧，为不臣，为更王；强国弱，小国强。”是时孙权受魏爵号，而称兵距守。◎弼按：孙权背魏，在黄初三年，与四年之太白昼见何涉？以是知天文、五行、符瑞诸志，多不足据也。】

            # Extra long string

            ''')
    res_list = result[1][0][1].as_list()
    print(res_list)
    print(''.join(flatten(res_list)))
    print(''.join(flatten(add_comment(res_list))))


def test_docx(pars):
    for par in pars:
        pf = par.paragraph_format
        if 1 or (pf.first_line_indent != None and int(pf.first_line_indent) > 0) or (pf.left_indent != None and int(pf.left_indent) > 0):
            print('{0} {4} -- {5} -- {1} -- {3} -- {2}'.format(pf.first_line_indent, pf.alignment, par.text[0:10], pf.space_before, pf.left_indent, par.style.name))


fn = FilenameInOut(args.file_name, dir_out=args.output_dir, ext_out='.Rmd')
doc = Document(fn.get_in_names()[0])

chars = pp.printables + ppu.Chinese.printables + '〇 　。，；：、﹑-！￥……（）―？《》〈〉〖〗■□♀．·“”„’‘◎○,;:/-!?<>~()[]{}#@$%^&*+=_/\\|`\'\'"'
# Combine() is space sensitive, see: https://stackoverflow.com/a/2940844
# content = pp.Combine(pp.Word(chars), adjacent = False)
content = pp.CharsNotIn('【】')
per_char = pp.Word(chars + '【】')
comment = pp.nested_expr('【', '】', content=content)
cmts_parser = ((content + comment[...]) ^ comment)[...]
pp.enable_all_warnings()

if args.test_cat == 'parser':
    test_cmts_parser(per_char)#cmts_parser)
    exit(0)
if args.test_cat == 'docx':
    test_docx(doc.paragraphs)
    exit(0)
    

# Parse nested comments as this post says: https://stackoverflow.com/a/5454510
parser = per_char if args.test_cat == 'char' else cmts_parser
styled_texts = []
FLAG_DEBUG = 0
chap_name = ''
sec_name = ''
i_sec_para = -1
for text in iter_text(doc.paragraphs):
    try:
        res = parser.parse_string(text, parse_all=True)
    except pp.ParseException as pe:
        print(pe.explain(depth=0))
    styled_text = ''.join(flatten(add_comment(res.as_list())))
    styled_texts.append(styled_text)

    if FLAG_DEBUG:
        print('[original text]: ' + text)
        print('[parsed text]: ' + styled_text)

attacher = MdIdAttacher('\n\n'.join([text.strip() for text in styled_texts if text.strip()]))

with open(fn.get_out_names()[0], 'w') as f:
    f.write(attacher.attached_full)

            